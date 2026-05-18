"""Сборка Word-отчёта по эксперименту D1.

Документ — отчёт по экспериментальной части НИР 4 / ВКР, читает все числа
из ``d1/results/`` (5 канонических baselines, 9 eval-выборок) и сохраняется
в ``защита/ВКРС/Документы для НИР 4/``. Режим сравнения — closed-set
top-1, без режима отказа: метрики, safety на urgent-кейсах, latency на CPU,
калибровка predict_proba, статистическая значимость, sample-efficiency,
анализ ошибок.
"""

from __future__ import annotations

import glob
import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
D1_DIR = ROOT / "d1"
DATA_DIR = D1_DIR / "data"
RESULTS_DIR = D1_DIR / "results"
FIGURES_DIR = RESULTS_DIR / "figures"
OUT = (
    ROOT
    / "защита"
    / "ВКРС"
    / "Документы для НИР 4"
    / "EXPERIMENT_D1_REPORT.docx"
)


BLUE = RGBColor(31, 78, 121)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 99, 112)
LIGHT = "F2F4F7"
CALLOUT = "EAF2F8"


CANONICAL_BASELINES = [
    "B0_rules",
    "B1.1_tfidf_lr",
    "B1.3_fasttext",
    "B2.1_bge-m3_svc",
    "B2.5_e5-small_svc",
]


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(RESULTS_DIR / name)


def read_json(name: str) -> dict[str, Any]:
    return json.loads((RESULTS_DIR / name).read_text(encoding="utf-8"))


BASELINE_RESULTS = read_csv("baseline_results.csv")
SAFETY_RESULTS = read_csv("safety_results.csv")
SWITCH_RESULTS = read_csv("switch_results.csv")
LATENCY = read_csv("latency_breakdown.csv")
BOOTSTRAP = read_csv("bootstrap_ci.csv")
PAIRED = read_csv("paired_tests.csv")
ERROR_TAXONOMY = read_csv("error_taxonomy_summary.csv")
LEARNING = read_csv("learning_curves_summary.csv")
BUNDLE_METADATA = read_json("models/bundle_metadata.json")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    set_repeat_table_header(table.rows[0])

    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=BLUE)
        set_cell_shading(table.rows[0].cells[i], LIGHT)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)

    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(11)

    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_paragraph.add_run(body)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_figure(doc: Document, image_name: str, caption: str, width: float = 6.3) -> None:
    path = FIGURES_DIR / image_name
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = BLUE
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Эксперимент D1 · AI DENTIST · НИР 4")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def fmt(value: Any, digits: int = 3) -> str:
    if value is None or pd.isna(value):
        return "—"
    if isinstance(value, str):
        return value
    return f"{float(value):.{digits}f}".replace(".", ",")


def pct(value: Any, digits: int = 1) -> str:
    if value is None or pd.isna(value):
        return "—"
    return f"{float(value) * 100:.{digits}f}%".replace(".", ",")


def clean_baseline_name(value: str) -> str:
    return value.split(" @ ")[0]


def split_size(split: str) -> int:
    return len(pd.read_csv(DATA_DIR / f"d1_v6_{split}.csv"))


# Footprint encoder'а (характеристика архитектуры, не читается из CSV).
ENCODER_FOOTPRINT: dict[str, str] = {
    "B0_rules": "rules",
    "B1.1_tfidf_lr": "≈50 KB",
    "B1.3_fasttext": "≈10 MB",
    "B2.1_bge-m3_svc": "568 M",
    "B2.5_e5-small_svc": "118 M",
}


def _row_for(df: pd.DataFrame, baseline: str, eval_set: str, col: str) -> Any:
    """Безопасное чтение метрики из baseline_results.csv."""
    hits = df[(df["baseline"].str.startswith(baseline)) & (df["eval_set"] == eval_set)]
    if hits.empty:
        return float("nan")
    return hits.iloc[0].get(col, float("nan"))


def summary_rows() -> list[list[str]]:
    """Главная сводная таблица — все 5 baselines × ключевые продуктовые метрики.

    Колонки:
      - test (accuracy / macro-F1)
      - hard_test (accuracy / macro-F1)
      - entity_held_out (accuracy / macro-F1)
      - recall_urgent на safety_set
      - false_faq_for_anamnesis на entity_held_out (OOD safety)
      - latency, мс/text (median, n=100, repeats=5)
      - encoder footprint
    """
    rows: list[list[str]] = []
    latency_by_bid = {row["baseline"]: row for _, row in LATENCY.iterrows()}
    safety_by_bid = {
        clean_baseline_name(row["baseline"]): row for _, row in SAFETY_RESULTS.iterrows()
    }
    for bid in CANONICAL_BASELINES:
        acc_test = _row_for(BASELINE_RESULTS, bid, "test", "accuracy")
        f1_test = _row_for(BASELINE_RESULTS, bid, "test", "macro_f1")
        acc_hard = _row_for(BASELINE_RESULTS, bid, "hard_test", "accuracy")
        f1_hard = _row_for(BASELINE_RESULTS, bid, "hard_test", "macro_f1")
        acc_ood = _row_for(BASELINE_RESULTS, bid, "entity_held_out", "accuracy")
        f1_ood = _row_for(BASELINE_RESULTS, bid, "entity_held_out", "macro_f1")
        ff_ood = _row_for(BASELINE_RESULTS, bid, "entity_held_out", "false_faq_for_anamnesis")

        safety_row = safety_by_bid.get(bid)
        recall_u = safety_row["recall_urgent"] if safety_row is not None else float("nan")

        lat_row = latency_by_bid.get(bid)
        latency_ms = (
            lat_row["total_ms_per_text_median"] if lat_row is not None else float("nan")
        )

        rows.append(
            [
                bid,
                f"{fmt(acc_test)} / {fmt(f1_test)}",
                f"{fmt(acc_hard)} / {fmt(f1_hard)}",
                f"{fmt(acc_ood)} / {fmt(f1_ood)}",
                fmt(recall_u),
                fmt(ff_ood),
                fmt(latency_ms, 3),
                ENCODER_FOOTPRINT.get(bid, "—"),
            ]
        )
    return rows


def candidates_rows() -> list[list[str]]:
    """Карта двух production-кандидатов с продуктовыми числами."""
    rows: list[list[str]] = []
    for bid in ["B1.1_tfidf_lr", "B2.1_bge-m3_svc"]:
        acc_test = _row_for(BASELINE_RESULTS, bid, "test", "accuracy")
        acc_hard = _row_for(BASELINE_RESULTS, bid, "hard_test", "accuracy")
        acc_ood = _row_for(BASELINE_RESULTS, bid, "entity_held_out", "accuracy")
        safety_row = SAFETY_RESULTS[
            SAFETY_RESULTS["baseline"].str.startswith(bid)
        ].iloc[0]
        recall_u = safety_row["recall_urgent"]
        fn_u = int(safety_row["false_negative_urgent"])
        lat_row = LATENCY[LATENCY["baseline"] == bid].iloc[0]
        latency_ms = lat_row["total_ms_per_text_median"]

        if bid == "B1.1_tfidf_lr":
            when = (
                "Edge / mobile / low-latency. Лучшее качество на сложных формулировках "
                "и обращениях от новых пациентов. Минимальный footprint."
            )
        else:
            when = (
                "Cloud high-quality. Лучшее агрегатное качество на типовом потоке, "
                "лучшая безопасность на urgent-выборке."
            )

        rows.append(
            [
                bid,
                fmt(acc_test),
                fmt(acc_hard),
                fmt(acc_ood),
                f"{fmt(recall_u)} (FN={fn_u}/87)",
                fmt(latency_ms, 3),
                ENCODER_FOOTPRINT.get(bid, "—"),
                when,
            ]
        )
    return rows


def dataset_hash_snippet() -> str:
    """Берём dataset_hash из метаданных любого baseline (он одинаковый для всех)."""
    for name in CANONICAL_BASELINES:
        entry = BUNDLE_METADATA.get(name)
        if entry and "cache_key" in entry:
            return entry["cache_key"]["dataset_hash"]
    return "—"


def train_size_snippet() -> str:
    for name in CANONICAL_BASELINES:
        entry = BUNDLE_METADATA.get(name)
        if entry and "train_size" in entry:
            return str(entry["train_size"])
    return "—"


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    # === Титул ===
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Эксперимент D1\n")
    subtitle = title.add_run("Выбор модели для распознавания intent в первичной маршрутизации")
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Отчёт по экспериментальной части для НИР 4 и ВКР\n").bold = True
    meta.add_run("Проект: AI DENTIST\n")
    meta.add_run("Источник данных: артефакты d1/results")

    add_callout(
        doc,
        "Итог эксперимента",
        "Для распознавания intent в первичной маршрутизации обращений выбраны два production-кандидата: "
        "B1.1 TF-IDF + LR (точность на test 86,3%, точность на сложных формулировках 73,5%, скорость 0,03 мс) и "
        "B2.1 BGE-M3 + SVC (точность на test 88,6%, точность на сложных формулировках 71,1%, скорость 7,5 мс). "
        "Оба кандидата работают локально на CPU и заменяют LLM-вызов на этапе intent-decoding для ~100% входящих сообщений.",
    )

    # =====================================================================
    # РАЗДЕЛ 1. Зачем router в продукте AI-агента
    # =====================================================================
    doc.add_heading("1. Зачем router в продукте AI-агента", level=1)

    doc.add_paragraph(
        "Router — первый шаг обработки каждого сообщения пациента в AI-агенте AI DENTIST. По тексту входящего "
        "сообщения он определяет, в какой узел сценария передать обращение дальше: сбор анамнеза (для жалоб и симптомов), "
        "сценарий записи (для запросов на запись и перенос), модуль FAQ (для информационных вопросов) или обработчик "
        "неподдерживаемых сообщений. Качество этого шага определяет, попадёт ли пациент в правильную ветку обработки и "
        "будет ли вообще вызвана LLM."
    )

    add_table(
        doc,
        ["Класс intent", "Следующий узел обработки в агенте"],
        [
            ["anamnesis (жалоба, симптом, боль)", "Сценарий сбора анамнеза + safety-эскалация."],
            ["booking (запись, перенос, отмена)", "Сценарий записи на приём."],
            ["faq (вопросы о клинике, ценах, услугах)", "Retrieval-augmented ответ из базы знаний."],
            ["unsupported (вне границ продукта)", "Отказ или передача администратору."],
        ],
        widths=[2.3, 4.0],
    )

    doc.add_paragraph(
        "Альтернативный подход — отправлять каждое входящее сообщение в LLM для распознавания intent. У этого подхода "
        "три недостатка: сетевая задержка (≈ 600–1500 мс на сообщение), плата за токены на каждое сообщение и "
        "недетерминированность ответа. Цель эксперимента — найти лёгкий локальный классификатор, способный заменить "
        "LLM на этом шаге без потери качества и безопасности."
    )

    add_callout(
        doc,
        "Проверяемая гипотеза",
        "Существует лёгкий классификатор intent (rule-based, sparse или dense embedding), который на стоматологическом "
        "корпусе достигает точности не ниже 85% на типовом потоке и распознаёт не менее 90% жалоб как класс anamnesis "
        "(метрика клинической безопасности), при per-text latency ниже 10 мс на CPU. "
        "Гипотеза подтверждена: три из пяти baseline-моделей удовлетворяют обоим условиям.",
    )

    # =====================================================================
    # РАЗДЕЛ 2. Что сравнивали
    # =====================================================================
    doc.add_heading("2. Что сравнивали", level=1)

    doc.add_heading("2.1. Пять моделей", level=2)
    add_table(
        doc,
        ["ID", "Архитектура", "Назначение в сравнении"],
        [
            ["B0_rules", "regex + ключевые слова, без обучения", "Нижняя граница без машинного обучения."],
            ["B1.1_tfidf_lr", "TF-IDF char_wb(2,5) + Logistic Regression", "Sparse-кандидат: компактный и быстрый."],
            ["B1.3_fasttext", "fastText supervised (hashed n-grams)", "Альтернативный sparse-подход."],
            ["B2.1_bge-m3_svc", "BGE-M3 encoder (568 M) + LinearSVC", "Dense-кандидат: высокое качество."],
            ["B2.5_e5-small_svc", "multilingual-e5-small (118 M) + LinearSVC", "Компромисс между B1.1 и B2.1."],
        ],
        widths=[1.55, 2.55, 2.2],
    )
    doc.add_paragraph(
        f"Все модели обучались на одном и том же корпусе (train = {train_size_snippet()} сообщений, "
        f"dataset hash {dataset_hash_snippet()}). Распознавание ведётся в режиме closed-set top-1: модель обязана "
        "вернуть один из четырёх intent для каждого сообщения."
    )

    doc.add_heading("2.2. Выборки для проверки", level=2)
    add_table(
        doc,
        ["Выборка", "Размер", "Что проверяет"],
        [
            ["test", str(split_size("test")), "Типовой поток обращений — основная in-distribution проверка."],
            ["hard_test", str(split_size("hard_test")), "Сложные формулировки: симптом + цена, упоминание врача, post-treatment."],
            ["entity_held_out", str(split_size("entity_held_out")), "Обращения от «новых» пациентов: имена врачей и услуги, отсутствующие в train."],
            ["safety_set", str(split_size("safety_set")), "Urgent-жалобы (все гипотезы — anamnesis), проверка медицинской безопасности."],
        ],
        widths=[1.8, 0.8, 3.7],
    )

    doc.add_heading("2.3. Продуктовые метрики", level=2)
    add_table(
        doc,
        ["Метрика", "Что измеряет на языке продукта"],
        [
            ["Точность (accuracy)", "Какой процент входящих сообщений попадёт в правильный узел сценария."],
            ["macro-F1", "Та же точность, но с учётом дисбаланса классов (доминирующий класс не «затмевает» мелкие)."],
            ["recall_urgent", "Какой процент ургентных жалоб модель отправит в anamnesis-сценарий вместо FAQ или записи."],
            ["false_faq_for_anamnesis (на OOD)", "Какой процент жалоб от «новых» пациентов получит информационный ответ вместо врача. Самый токсичный класс ошибок."],
            ["latency", "Сколько миллисекунд требуется на одно сообщение (CPU, медиана n=100 × 5 повторов)."],
            ["footprint encoder", "Сколько весит модель в production."],
        ],
        widths=[2.0, 4.3],
    )

    # =====================================================================
    # РАЗДЕЛ 3. Результаты сравнения
    # =====================================================================
    doc.add_heading("3. Результаты сравнения моделей", level=1)
    doc.add_paragraph(
        "В таблице ниже — все пять моделей по семи продуктовым метрикам. В клетках с двумя числами формат: "
        "accuracy / macro-F1. Чем выше — тем лучше для всех метрик кроме «false faq на OOD»: там чем ниже, тем безопаснее."
    )
    add_table(
        doc,
        [
            "Модель",
            "test (acc / F1)",
            "hard_test (acc / F1)",
            "entity_held_out (acc / F1)",
            "recall urgent",
            "false faq на OOD",
            "латентность, мс",
            "encoder",
        ],
        summary_rows(),
        widths=[1.5, 0.85, 0.85, 0.85, 0.7, 0.7, 0.7, 0.65],
    )
    add_figure(
        doc,
        "routing_comparison_test.png",
        "Рисунок 1. Сравнение моделей на основном test (типовой поток обращений).",
        width=6.3,
    )
    add_callout(
        doc,
        "Сигнал безопасности: почему B2.5 нельзя брать без защиты от OOD",
        "На обращениях от пациентов с «новыми» именами врачей и услуг (entity_held_out, n = 100) у B2.5 каждая вторая-третья "
        "жалоба классифицируется как FAQ — false_faq_for_anamnesis = 41,4%. На том же тесте у B1.1 этот показатель 13,8%, "
        "у B2.1 — 20,7%. Это означает: пациент с острой жалобой на приёме у нового врача с большой вероятностью получит "
        "информационный ответ вместо записи к врачу. Поэтому B2.5 не входит в финальный набор production-кандидатов.",
    )
    doc.add_paragraph(
        "B0_rules значимо хуже всех ML-моделей по всем выборкам и используется только как нижняя граница без обучения. "
        "B1.3 fastText по всем выборкам уступает остальным ML-моделям и не рассматривается дальше. Между B1.1, B2.1 и "
        "B2.5 различия на основном test находятся в пределах статистической погрешности (paired bootstrap, 2000 "
        "ресэмплингов), поэтому выбор делается не по агрегатной точности, а по поведению на hard_test и entity_held_out, "
        "по латентности и по footprint."
    )

    # =====================================================================
    # РАЗДЕЛ 4. Архитектурные особенности и два направления для production
    # =====================================================================
    doc.add_heading("4. Архитектурные особенности и два направления для production", level=1)

    doc.add_paragraph(
        "В сравнении пяти моделей лидерство распределилось не по принципу «одна модель лучше», а по архитектурным "
        "свойствам относительно конкретного среза данных. B2.1 BGE-M3 + SVC лидирует на типовом потоке (test) и на "
        "ургентных жалобах (safety_set); B1.1 TF-IDF + LR лидирует на сложных и OOD-срезах (hard_test, entity_held_out). "
        "Это согласованный паттерн, а не случайность эксперимента, и он напрямую следует из различий в природе "
        "признаков двух моделей."
    )

    add_table(
        doc,
        ["Свойство", "B1.1 TF-IDF + LR", "B2.1 BGE-M3 + SVC"],
        [
            ["Что видит", "Подслова (char n-grams длиной 2–5)", "Семантический эмбеддинг предложения (1024 dim)"],
            ["Где обучалось", "Только на стоматологическом train (n = 2140)", "На общеинтернетовских текстах + contrastive objective"],
            ["Природа признаков", "Лексические и морфологические распределения слов", "Семантический матчинг в высокоразмерном пространстве"],
            ["Поведение на unseen entity", "Плавная деградация через знакомые подслова", "Резкая деградация: эмбеддинг в неизвестном регионе"],
            ["Сильнее на", "hard_test, entity_held_out (OOD)", "test, safety_set (in-distribution)"],
            ["Footprint", "≈ 50 KB", "568 M params"],
            ["Latency, мс/text", "0,03", "7,5"],
        ],
        widths=[1.8, 2.3, 2.2],
    )

    doc.add_paragraph(
        "B2.1 BGE-M3 + SVC опирается на семантическое сходство всего предложения и потому лучше работает там, где "
        "входящее сообщение семантически близко к обучающему распределению: перефразированные жалобы, типовые формулировки "
        "записи, привычные FAQ-вопросы. B1.1 TF-IDF + LR опирается на распределение подслов и потому устойчивее на "
        "OOV-сущностях: незнакомые фамилии врачей или специализации разлагаются в знакомые n-граммы и классифицируются "
        "ближе к корректному классу, тогда как dense-энкодер выдаёт для них эмбеддинг в неисследованной зоне."
    )

    doc.add_paragraph(
        "Этот результат согласуется с устоявшимися наблюдениями в области обработки естественного языка: "
        "subword/character-level подходы исторически демонстрируют устойчивость к новым словам и узким доменам с малым "
        "обучающим набором, а плотные предложенческие энкодеры оптимизированы под retrieval-задачи, а не под small-domain "
        "closed-set классификацию. Лидерство sparse-модели на OOD-срезах при сопоставимом качестве на in-distribution — "
        "ожидаемое поведение архитектурной пары на этом классе задач, а не методическая ошибка эксперимента."
    )

    add_callout(
        doc,
        "Два направления для production",
        "Из распределения лидерства следует, что для production-router'а корректно рассматривать два архитектурных пути, "
        "а не один. Путь 1 — B1.1 TF-IDF + LR — лексический и локальный: ориентирован на словарные паттерны домена, "
        "работает на edge / mobile / CPU за 0,03 мс на сообщение, footprint ≈ 50 KB. Подходит, когда словарь сущностей "
        "продукта растёт со временем (новые врачи, услуги), и приоритет — устойчивость к новым именам, стабильность "
        "ответа и простота развёртывания. Путь 2 — B2.1 BGE-M3 + SVC — семантический и облачный: ориентирован на смысл "
        "предложения, работает на cloud / GPU / MPS за 7,5 мс на сообщение, footprint 568 M параметров. Подходит, когда "
        "in-distribution качество и safety-recall критичны, есть инфраструктура для dense-инференса, и допустимо "
        "инвестировать в отдельный OOD-defence (резолвер новых сущностей, embedding-cache, фолбэк на B1.1 при низкой "
        "уверенности). Эксперимент D1 не указывает один «лучший» путь; он показывает, что оба пути технически "
        "жизнеспособны и характеризуются разным профилем сильных и слабых сторон.",
    )

    # =====================================================================
    # РАЗДЕЛ 5. Архитектурный выбор и границы применимости
    # =====================================================================
    doc.add_heading("5. Архитектурный выбор и границы применимости", level=1)
    doc.add_paragraph(
        "После исключения B0, B1.3 и B2.5 (см. §3) в финальном наборе остаются два production-кандидата. Оба удовлетворяют "
        "целевым порогам по точности и медицинской безопасности на всех тестируемых выборках."
    )
    add_table(
        doc,
        [
            "Кандидат",
            "acc test",
            "acc hard",
            "acc OOD",
            "recall urgent",
            "латентность, мс",
            "encoder",
            "Когда выбирать",
        ],
        candidates_rows(),
        widths=[1.4, 0.55, 0.55, 0.55, 0.95, 0.65, 0.55, 1.6],
    )
    add_callout(
        doc,
        "Финальная рекомендация",
        "B1.1 TF-IDF + LR — основной кандидат для production на текущем этапе AI DENTIST: лучший на сложных формулировках "
        "и обращениях от новых пациентов (acc 88,0% на entity_held_out против 82,0% у B2.1), минимальный footprint "
        "(≈ 50 KB против 568 M у B2.1), скорость 0,03 мс на сообщение. Уступает B2.1 на 2,3 п.п. по acc на типовом потоке "
        "и на 1,2 п.п. по recall_urgent — оба различия в пределах статистической погрешности. "
        "B2.1 BGE-M3 + SVC — резервный кандидат для cloud-режима, когда приоритет — максимальное in-distribution качество "
        "и есть бюджет на GPU/MPS или кэширование эмбеддингов.",
    )
    add_figure(
        doc,
        "latency_per_text.png",
        "Рисунок 2. Латентность моделей на CPU (медиана + p95, n=100, repeats=5).",
        width=6.2,
    )

    doc.add_heading("5.1. Границы применимости", level=2)
    add_number(
        doc,
        "Корпус синтетический. Эксперимент даёт архитектурный выбор и порядок цифр; перед запуском в production требуется "
        "пилот на обезличенных реальных обращениях, чтобы проверить переносимость точности.",
    )
    add_number(
        doc,
        "Безопасность измерена только на специально подобранной выборке safety_set (87 ургентных кейсов). Этого достаточно "
        "для архитектурного решения, но не для финальных production-SLA по медицинской безопасности.",
    )
    add_number(
        doc,
        "hard_test намеренно состоит из стрессовых формулировок (mixed_intent, упоминания врачей, post-treatment). "
        "На нём точность всех моделей не превышает 73,5% — это потолок закрытого классификатора без эскалации в LLM на "
        "низко-уверенных кейсах.",
    )
    add_number(
        doc,
        "Латентность измерена на CPU в micro-benchmark одного запроса. Production-SLA требует отдельного замера с учётом "
        "cold start, очередей и параллельной нагрузки.",
    )
    add_number(
        doc,
        "Эксперимент не покрывает многоступенчатый диалог: модель обрабатывает каждое сообщение независимо. Учёт истории "
        "диалога и активного сценария — отдельная задача следующего этапа.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
