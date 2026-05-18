from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUN_DIR = ROOT / "d4" / "outputs" / "runs" / "20260419_235319"
OUT = (
    ROOT
    / "защита"
    / "ВКРС"
    / "Документы для НИР 4"
    / "EXPERIMENT_D4_REPORT.docx"
)


BLUE = RGBColor(31, 78, 121)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 99, 112)
LIGHT = "F2F4F7"
CALLOUT = "EAF2F8"


def load_report(mode: str) -> dict:
    return json.loads((RUN_DIR / f"full_dev_{mode}.report.json").read_text(encoding="utf-8"))


REPORTS = {mode: load_report(mode) for mode in ["plain", "contextual", "llm_enriched"]}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=BLUE)
        set_cell_shading(table.rows[0].cells[i], LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = BLUE
    r.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


def add_figure(doc: Document, image_name: str, caption: str, width: float = 6.3) -> None:
    path = RUN_DIR / "figures" / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = BLUE
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Эксперимент D4 · AI DENTIST · НИР 4")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def pct(value: float) -> str:
    return f"{value * 100:.1f}%".replace(".", ",")


def f4(value: float) -> str:
    return f"{value:.4f}".replace(".", ",")


def build_doc() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    # Cover page
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Эксперимент D4\n")
    subtitle = title.add_run("Выбор стратегии работы с базой знаний клиники")
    subtitle.font.size = Pt(16)
    subtitle.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Отчёт по экспериментальной части для НИР 4 и ВКРС\n").bold = True
    meta.add_run("Проект: AI DENTIST\n")
    meta.add_run("Канонический прогон: 20260419_235319 · полный валидационный набор, 137 вопросов\n")
    meta.add_run("Дата подготовки: 10 мая 2026 года")

    add_callout(
        doc,
        "Краткий вывод",
        "На текущей базе одной стоматологической клиники лучшим проверенным решением по соотношению качества, скорости и сложности оказался простой векторный поиск по обычным карточкам знаний. Усложнение поиска и поверхностное обогащение карточек не дали устойчивого выигрыша, достаточного для включения в промышленный контур MVP.",
    )

    doc.add_page_break()

    doc.add_heading("1. Назначение документа", level=1)
    doc.add_paragraph(
        "Документ фиксирует ход, методику, результаты и практическое значение эксперимента D4. "
        "Он не является полной главой ВКРС, а служит опорным отчётом по конкретному эксперименту: "
        "какая гипотеза проверялась, на каких данных, какими методами, какие результаты получены "
        "и как они влияют на архитектуру проекта AI DENTIST."
    )
    add_bullet(doc, "Объект эксперимента: модуль ответов на частые вопросы по базе знаний стоматологической клиники.")
    add_bullet(doc, "Предмет эксперимента: способы поиска релевантных карточек знаний и способы представления карточек перед индексированием.")
    add_bullet(doc, "Прикладная задача: выбрать рациональный контур доступа к знаниям для MVP без неоправданного усложнения системы.")

    doc.add_heading("2. Связь эксперимента с проектом AI DENTIST", level=1)
    doc.add_paragraph(
        "AI DENTIST разрабатывается как модульная система первичной текстовой коммуникации стоматологической клиники с пациентом. "
        "В этой системе модуль ответов на частые вопросы отвечает за информационные запросы по данным конкретной клиники: "
        "врачи, цены, режим работы, документы, анестезия, рекомендации после процедур и организационные правила."
    )
    doc.add_paragraph(
        "Эксперимент D4 нужен для того, чтобы не выбирать архитектуру поиска интуитивно. "
        "В медицинском домене лишняя сложность опасна: каждый дополнительный компонент увеличивает задержку ответа, стоимость, число точек отказа и риск труднообъяснимой ошибки. "
        "Поэтому эксперимент проверяет, нужен ли тяжёлый поисковый контур или для текущего масштаба базы достаточно более простой схемы."
    )

    add_table(
        doc,
        ["Связь с продуктом", "Что даёт эксперимент"],
        [
            ["Качество ответа", "Проверяет, попадают ли обязательные факты из базы знаний в ответ пациенту."],
            ["Безопасность", "Отдельно учитывает правильный отказ на вопросы вне базы знаний."],
            ["Скорость", "Позволяет отказаться от лишних этапов, если они не дают выигрыша."],
            ["Стоимость", "Снижает число обращений к дополнительным моделям и стоимость обновления базы знаний."],
            ["Архитектура", "Определяет, какой контур включать в MVP, а какие ветки оставить исследовательскими."],
        ],
        widths=[2.1, 4.2],
    )

    doc.add_heading("3. Постановка исследовательской задачи", level=1)
    doc.add_paragraph(
        "Главный исследовательский вопрос: какой способ поиска ответа в базе знаний клиники и какое представление карточек знаний дают наилучшее соотношение качества ответа и сложности реализации?"
    )
    doc.add_heading("3.1. Гипотеза этапа 1", level=2)
    doc.add_paragraph(
        "Усложнение способа поиска — переход от простого векторного поиска к гибридному поиску, переранжированию или каскадной маршрутизации — даёт измеримый и устойчивый прирост качества ответа."
    )
    doc.add_heading("3.2. Гипотеза этапа 2А", level=2)
    doc.add_paragraph(
        "Обогащение карточек знаний структурным заголовком или коротким пересказом от языковой модели улучшает поиск и итоговый ответ без смены самой механики поиска."
    )
    doc.add_paragraph(
        "Обе гипотезы проверяют практическую интуицию: более сложный контур должен быть оправдан только тогда, когда он даёт измеримый выигрыш на данных."
    )

    doc.add_heading("4. Данные и контролируемые условия", level=1)
    add_table(
        doc,
        ["Компонент данных", "Объём"],
        [
            ["Карточки знаний для индекса", "63"],
            ["Карточки врачей", "22"],
            ["Полный валидационный набор", "137 вопросов"],
            ["Вопросы с эталонными карточками", "121 из 137"],
            ["Вопросы с обязательным врачом", "23 из 137"],
            ["Тестовый набор", "51 вопрос; на текущей итерации не использовался"],
        ],
        widths=[3.5, 2.8],
    )
    doc.add_paragraph(
        "База знаний была зафиксирована до полного прогона. Разметка эталонных карточек и обязательных фактов проходила аудит по единой политике аннотации. "
        "Это снижает риск того, что результаты объясняются изменением данных между прогонами."
    )

    doc.add_heading("5. Использованные модели и роли", level=1)
    add_table(
        doc,
        ["Роль", "Модель", "Использование"],
        [
            ["Генератор ответа", "Qwen/Qwen3-Next-80B-A3B-Instruct", "Все стратегии с языковой моделью; также обогащение карточек C2."],
            ["Модель-судья", "openai/gpt-5.4-mini", "Оценка доли обязательных фактов, попавших в ответ."],
            ["Модель векторных представлений", "BAAI/bge-m3", "Векторный поиск S3, гибридный поиск S4, исследовательские ветки S4r и S5."],
            ["Модель переранжирования", "BAAI/bge-reranker-base", "Только исследовательские ветки S4r и S5."],
            ["Токенизатор", "Qwen/Qwen3.5-35B-A3B", "Подсчёт длины карточек и контекста."],
        ],
        widths=[1.55, 2.25, 2.5],
    )

    doc.add_heading("6. Сравниваемые подходы", level=1)
    add_table(
        doc,
        ["Обозначение", "Русское описание", "Смысл"],
        [
            ["B0", "Правила и шаблоны", "Быстрые ответы по ключевым словам без языковой модели."],
            ["S1", "Подача всей базы знаний", "Вся база передаётся модели как контекст; используется как ориентир качества, но плохо масштабируется."],
            ["S2", "Поиск по словам", "Ищет совпадения терминов и формулировок."],
            ["S3", "Простой векторный поиск", "Ищет карточки по смысловой близости вопроса и текста карточки."],
            ["S4", "Гибридный поиск", "Объединяет поиск по словам и смысловой поиск."],
            ["S4r", "Гибридный поиск с переранжированием", "После первичного поиска отдельная модель пересортировывает кандидатов."],
            ["S5", "Каскад правил и гибрида", "Сначала пробует шаблонные ответы, затем переходит к гибриду с переранжированием."],
        ],
        widths=[0.75, 1.85, 3.7],
    )
    add_table(
        doc,
        ["Представление", "Описание", "Ожидаемый эффект"],
        [
            ["C0", "Обычная карточка", "Контрольный вариант без обогащения."],
            ["C1", "Структурный заголовок", "Должен явно сообщать раздел и тип карточки."],
            ["C2", "Короткий пересказ от модели", "Должен помочь поиску за счёт более явного описания смысла карточки."],
        ],
        widths=[1.0, 2.5, 2.8],
    )

    doc.add_heading("7. Метрики качества", level=1)
    add_table(
        doc,
        ["Метрика", "Русское пояснение", "Почему важна"],
        [
            ["hit@5", "Доля вопросов, где хотя бы одна правильная карточка попала в первые пять результатов.", "Показывает, видит ли модель нужный источник."],
            ["MRR", "Средняя обратная позиция первой правильной карточки.", "Показывает, насколько высоко находится правильный источник."],
            ["mean_recall", "Средняя доля всех правильных карточек, найденных в первой пятёрке.", "Важна для вопросов, где нужно несколько источников."],
            ["answerability_correct", "Правильное решение отвечать или отказать.", "Критично для безопасности: нельзя отвечать вне базы знаний."],
            ["doctor_match", "Правильное указание врача на релевантных вопросах.", "Важно для маршрутизации пациента."],
            ["FMR", "Доля обязательных фактов из эталонной разметки, попавших в ответ.", "Главная содержательная метрика качества ответа."],
        ],
        widths=[1.35, 3.0, 2.0],
    )

    doc.add_heading("8. Ход эксперимента", level=1)
    add_number(doc, "Малый стрессовый набор: проверка инфраструктуры и различимости стратегий.")
    add_number(doc, "Пилотный прогон: предварительная проверка поисковых стратегий и отсечение сложных веток без явного сигнала.")
    add_number(doc, "Полный валидационный прогон: основной результат на 137 вопросах.")
    add_number(doc, "Проверка представлений карточек: сравнение C0, C1 и C2 при фиксированном поисковом контуре.")
    add_number(doc, "Статистическая проверка различий по доступным посегментным метрикам.")

    doc.add_heading("9. Результаты этапа 1: выбор способа поиска", level=1)
    plain = REPORTS["plain"]
    retrieval_rows = []
    for sid in ["S2", "S3", "S4"]:
        r = plain["retrieval"][sid]
        retrieval_rows.append([sid, pct(r["hit_rate"]), f4(r["mrr"]), f4(r["mean_recall"]), pct(r["gold_in_context_rate"])])
    add_table(
        doc,
        ["Стратегия", "Попадание в первую пятёрку", "Средняя позиция", "Средняя полнота", "Полный эталон в контексте"],
        retrieval_rows,
        widths=[0.8, 1.35, 1.2, 1.2, 1.55],
    )
    quality_rows = []
    names = {
        "B0": "Правила",
        "S1": "Вся база",
        "S2": "Слова",
        "S3": "Вектор",
        "S4": "Гибрид",
    }
    for sid in ["B0", "S1", "S2", "S3", "S4"]:
        d = plain["deterministic"][sid]
        quality_rows.append(
            [
                f"{sid} — {names[sid]}",
                f'{d["answerability_correct"]}/137',
                f'{d["doctor_match"]}/{d["doctor_total"]}',
                f4(d["avg_fmr"]),
            ]
        )
    add_table(
        doc,
        ["Стратегия", "Правильное решение", "Правильный врач", "Доля обязательных фактов"],
        quality_rows,
        widths=[2.0, 1.3, 1.2, 1.6],
    )
    add_callout(
        doc,
        "Вывод этапа 1",
        "На полном валидационном наборе простой векторный поиск S3 оказался лучшим проверенным поисковым контуром среди S2, S3 и S4. Он дал 95,0% попадания правильной карточки в первую пятёрку, 134 правильных решения отвечать или отказать из 137 и 0,6190 по доле обязательных фактов. Гибридный поиск S4 не улучшил результат, а исследовательские ветки S4r и S5 не были вынесены в полный прогон после пилота.",
    )
    add_figure(doc, "stage1_retrieval_strategies.png", "Рисунок 1. Сравнение поисковых стратегий по метрикам поиска.", width=6.4)
    add_figure(doc, "stage1_quality_strategies.png", "Рисунок 2. Сравнение стратегий по качеству итогового ответа.", width=6.4)

    doc.add_heading("10. Результаты этапа 2А: представление карточек знаний", level=1)
    stage2_rows = []
    for mode, label in [
        ("plain", "C0 — обычное представление"),
        ("contextual", "C1 — структурный заголовок"),
        ("llm_enriched", "C2 — пересказ от модели"),
    ]:
        r = REPORTS[mode]["retrieval"]["S3"]
        d = REPORTS[mode]["deterministic"]["S3"]
        stage2_rows.append(
            [
                label,
                pct(r["hit_rate"]),
                f4(r["mrr"]),
                f'{d["answerability_correct"]}/137',
                f'{d["doctor_match"]}/{d["doctor_total"]}',
                f4(d["avg_fmr"]),
            ]
        )
    add_table(
        doc,
        ["Представление", "Попадание в первую пятёрку", "Средняя позиция", "Правильное решение", "Врач", "Доля фактов"],
        stage2_rows,
        widths=[2.0, 1.1, 1.0, 1.1, 0.8, 1.0],
    )
    doc.add_paragraph(
        "Структурный заголовок C1 практически не изменил качество. Пересказ от модели C2 немного повысил долю обязательных фактов по агрегированной оценке, но одновременно снизил правильность решения отвечать или отказать: 129 из 137 против 134 из 137 у обычного представления. По основной метрике FMR нет посегментной статистической проверки, поэтому небольшой агрегированный прирост не считается достаточным основанием для внедрения."
    )
    add_callout(
        doc,
        "Вывод этапа 2А",
        "Поверхностное обогащение карточек не дало устойчивого выигрыша. Для MVP сохраняется обычное представление карточек C0. Следующая исследовательская задача должна быть связана не с приклеиванием текста к карточкам, а с границами карточек, структурой фактов и проверкой ответа.",
    )
    add_figure(doc, "retrieval_comparison.png", "Рисунок 3. Сравнение представлений карточек по метрикам поиска.", width=6.4)
    add_figure(doc, "quality_comparison.png", "Рисунок 4. Сравнение представлений карточек по качеству ответа.", width=6.4)
    add_figure(doc, "mode_deltas.png", "Рисунок 5. Изменения метрик относительно обычного представления карточек.", width=6.2)
    add_figure(doc, "rank_shift.png", "Рисунок 6. Изменение позиций эталонных карточек при разных представлениях.", width=6.2)

    doc.add_heading("11. Проверка случайности различий", level=1)
    add_table(
        doc,
        ["Метрика", "C1 относительно C0", "C2 относительно C0", "Интерпретация"],
        [
            ["MRR", "−0,0138; p = 0,303", "−0,0092; p = 0,557", "Различия в пределах случайности."],
            ["hit@5", "−0,0165; p = 0,625", "−0,0248; p = 0,453", "Нет подтверждённого выигрыша."],
            ["Правильное решение", "−0,0146; p = 0,500", "−0,0365; p = 0,062", "Для C2 есть направленный отрицательный сигнал, но строгий порог значимости не пройден."],
            ["Правильный врач", "+0,0000; p = 1,000", "+0,0435; p = 1,000", "Выборка мала: 23 релевантных вопроса."],
        ],
        widths=[1.25, 1.45, 1.45, 2.2],
    )
    doc.add_paragraph(
        "Посегментная проверка по доле обязательных фактов невозможна в текущей версии конвейера, потому что оценки модели-судьи сохраняются только агрегированно. Это ограничение отдельно учитывается в выводах."
    )

    doc.add_heading("12. Итоговое архитектурное решение", level=1)
    add_callout(
        doc,
        "Рекомендуемая конфигурация",
        "Для текущего MVP рекомендуется простой векторный поиск S3 по обычным карточкам C0. Эта конфигурация сохраняет высокое качество поиска, не требует дополнительной модели переранжирования и не требует обогащения карточек при обновлении базы знаний.",
    )
    add_table(
        doc,
        ["Компонент", "Решение"],
        [
            ["Поисковый контур", "S3 — простой векторный поиск"],
            ["Представление карточек", "C0 — обычное представление"],
            ["Число карточек в контексте", "5"],
            ["Модель ответа", "Qwen/Qwen3-Next-80B-A3B-Instruct"],
            ["Модель векторных представлений", "BAAI/bge-m3"],
            ["Что не включается в MVP", "Гибрид, переранжирование, каскад и обогащение карточек как обязательный путь."],
        ],
        widths=[2.2, 4.0],
    )

    doc.add_heading("13. Влияние эксперимента на проект", level=1)
    doc.add_paragraph(
        "Эксперимент D4 влияет на проект не как самостоятельная демонстрация алгоритма, а как архитектурный фильтр. "
        "Он показывает, какие элементы не следует включать в промышленный контур без дополнительного доказательства пользы."
    )
    add_bullet(doc, "Модуль ответов по базе клиники остаётся отдельным управляемым узлом внутри общей архитектуры AI DENTIST.")
    add_bullet(doc, "Для MVP выбирается быстрый и простой путь доступа к знаниям, что снижает стоимость и задержку ответа.")
    add_bullet(doc, "Переранжирование и каскад не удаляются как идея, но переводятся в исследовательский или условный контур для сложных вопросов.")
    add_bullet(doc, "Следующий прирост качества ожидается не от поверхностного обогащения карточек, а от лучшей структуры базы знаний, более точного разбиения карточек и проверки ответа.")
    add_bullet(doc, "Для медицинского сценария сохраняется необходимость отказа на вопросы вне базы и передачи сложных случаев человеку.")

    doc.add_heading("14. Ограничения", level=1)
    add_number(doc, "Эксперимент проведён на базе одной стоматологической клиники; перенос на другие клиники требует повторной проверки.")
    add_number(doc, "Размер базы знаний мал: 63 карточки. На большой базе эффект гибридного поиска и переранжирования может измениться.")
    add_number(doc, "Гибрид с переранжированием S4r и каскад S5 не проверялись на полном валидационном наборе.")
    add_number(doc, "Тестовый и слепой наборы не использовались в каноническом полном прогоне.")
    add_number(doc, "Доля обязательных фактов сохраняется агрегированно, поэтому для неё пока нет строгой парной проверки по каждому вопросу.")
    add_number(doc, "Метрика безосновных утверждений используется только как контрольная, так как текущий парсер не ловит все типы фактических ошибок.")

    doc.add_heading("15. Что требуется проверить дальше", level=1)
    add_bullet(doc, "Прогнать независимый тестовый и слепой наборы после фиксации конфигурации.")
    add_bullet(doc, "Сохранять долю обязательных фактов по каждому вопросу для полноценной статистической проверки.")
    add_bullet(doc, "Проверить условное переранжирование только для сложных и рискованных вопросов, а не для каждого запроса.")
    add_bullet(doc, "Провести отдельный эксперимент по границам карточек: один врач, одна услуга, одна процедура или один блок рекомендаций как отдельная единица поиска.")
    add_bullet(doc, "Подключить реальные вопросы из пилота клиники и сравнить их с синтетическим валидационным набором.")

    doc.add_heading("16. Заключение", level=1)
    doc.add_paragraph(
        "Эксперимент D4 показал, что для текущей небольшой базы знаний стоматологической клиники усложнение поискового контура не получило достаточного экспериментального подтверждения. "
        "Простой векторный поиск по обычным карточкам знаний оказался наиболее рациональным решением для MVP: он обеспечивает высокое качество поиска, хорошую корректность решения отвечать или отказать и не требует дополнительных моделей в обязательном пути."
    )
    doc.add_paragraph(
        "Главный результат эксперимента состоит не в универсальном доказательстве превосходства одной поисковой стратегии, а в локализации узкого места. "
        "Дальнейшее развитие модуля должно быть направлено на структуру базы знаний, качество карточек, проверку ответа и безопасную маршрутизацию, а не на механическое добавление более сложных поисковых компонентов."
    )

    doc.add_heading("17. Источники", level=1)
    sources = [
        "Lewis P. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020. URL: https://papers.neurips.cc/paper_files/paper/2020/hash/6b493230205f780e1bc26945df7481e5-Abstract.html",
        "Chen J. et al. BGE M3-Embedding: Multi-Lingual, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation. 2024. URL: https://huggingface.co/papers/2402.03216",
        "Anthropic Engineering. Contextual Retrieval in AI Systems. 2024. URL: https://www.anthropic.com/engineering/contextual-retrieval",
        "Saad-Falcon J. et al. ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems. NAACL, 2024. URL: https://aclanthology.org/2024.naacl-long.20/",
        "Es S. et al. RAGAS: Automated Evaluation of Retrieval Augmented Generation. EACL, 2024. URL: https://aclanthology.org/2024.eacl-demo.16/",
        "Peng X. et al. Unanswerability Evaluation for Retrieval Augmented Generation. 2024. URL: https://arxiv.org/abs/2412.12300",
        "Zhang G. et al. Leveraging long context in retrieval augmented language models for medical question answering. npj Digital Medicine, 2025. URL: https://www.nature.com/articles/s41746-025-01651-w",
        "Yang Q. et al. Dual retrieving and ranking medical large language model with retrieval augmented generation. Scientific Reports, 2025. URL: https://www.nature.com/articles/s41598-025-00724-w",
        "Muhetaer M. et al. Medical QA dialogue datasets in RAG systems performance evaluation and ChatGPT optimization. Scientific Reports, 2025. URL: https://www.nature.com/articles/s41598-025-28015-4",
    ]
    for s in sources:
        add_number(doc, s)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Приложение. Канонические артефакты", level=1)
    add_table(
        doc,
        ["Артефакт", "Назначение"],
        [
            ["d4/EXPERIMENT_D4.md", "Единый источник правды по эксперименту D4."],
            ["d4/outputs/runs/20260419_235319/full_dev_plain.report.json", "Основные метрики для обычного представления карточек."],
            ["d4/outputs/runs/20260419_235319/full_dev_contextual.report.json", "Метрики для структурного заголовка."],
            ["d4/outputs/runs/20260419_235319/full_dev_llm_enriched.report.json", "Метрики для обогащения пересказом."],
            ["d4/outputs/runs/20260419_235319/reports/stage1_decision_memo.md", "Автоматически сформированный вывод по этапу 1."],
            ["d4/outputs/runs/20260419_235319/reports/stage2a_decision_memo.md", "Автоматически сформированный вывод по этапу 2А."],
        ],
        widths=[3.4, 2.9],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
